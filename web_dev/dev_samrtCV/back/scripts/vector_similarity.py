import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from pdfminer.high_level import extract_text as extract_text_pdf
import spacy
import pandas as pd
import re
from datetime import datetime, timedelta
import locale
from unidecode import unidecode
from sentence_transformers import SentenceTransformer
import torch.nn.functional as F

# Load the language model
nlp = spacy.load('fr_core_news_sm')

def traitement(cv_folder_path, job_offre_path):
    def extract_text_from_pdf(pdf_path):
        return extract_text_pdf(pdf_path)

    def extract_text_from_docx(docx_path):
        doc = Document(docx_path)
        return ' '.join([p.text for p in doc.paragraphs])

    def preprocess_text(text):
        # Implement your text preprocessing here (lowercasing, punctuation removal, etc.)
        return text

    # Assuming 'cv_folder_path' is the path to the folder containing CVs in PDF format
    # and 'job_offer_path' is the path to the job offer in DOCX format.

    cv_folder_path = r'C:\Users\guedrouz.ESH\Downloads\web_dev\dev_samrtCV\back\uploads\cv' # Example path to the folder containing CVs
    job_offer_path = r'C:\Users\guedrouz.ESH\Downloads\web_dev\dev_samrtCV\back\uploads\jobs\Data_Analyst.docx'

    # Step 1: Extract text from the files
    cv_info_list = []

    # Extract text from job offer
    """with open(job_offer_path, 'r', encoding='latin-1') as file:
        job_offer_text = file.read()"""
    def extract_text_from_docx(docx_path):
        doc = Document(docx_path)
        return ' '.join([p.text for p in doc.paragraphs])
    
    job_offer_text = extract_text_from_docx(job_offer_path)
    job_offer_text = preprocess_text(job_offer_text)

    #print(job_offer_text)
    
    competences = [
    # Compétences en Data Science
    'python',
    'r',
    'julia',
    'scala',
    'pandas',
    'numpy',
    'scikit-learn',
    'tensorflow',
    'pytorch',
    'keras',
    'matplotlib',
    'seaborn',
    'plotly',
    'dask',
    'spark',
    'sql',
    'mongodb',
    'cassandra',
    'couchdb',
    'bigquery',
    'aws redshift',
    'snowflake',
    'hadoop',
    'apache spark',
    'apache flink',
    'data cleaning',
    'data wrangling',
    'feature engineering',
    'regression linéaire',
    'regression logistique',
    'tests d\'hypothèses',
    'analyse de variance',
    'analyse de séries temporelles',
    'classification',
    'régression',
    'clustering',
    'random forest',
    'gradient boosting',
    'svm',
    'réseaux de neurones',
    'réseaux de neurones convolutifs',
    'réseaux de neurones récurrents',
    'autoencodeurs',
    'transformers',
    'tokenization',
    'word embeddings',
    'modèles de langage',
    'bert',
    'gpt',
    'ner',
    'sentiment analysis',
    'tableau',
    'power bi',
    'flask',
    'fastapi',
    'docker',
    'kubernetes',
    'aws sagemaker',
    'tensorflow serving',
    'création de nouvelles caractéristiques',
    'sélection de caractéristiques',
    'traitement des valeurs manquantes',
    'normalisation',
    'standardisation',
    'méthodologies agiles',
    'git',
    'jira',
    'trello',
    'communication des résultats',
    'rédaction de rapports techniques',
    'présentations',
    'compréhension des besoins métier',
    "statistique",
    "machine learning",
    "deep learning",
    "data",
    "capacité d'analyse",
    "force de proposition",
    "bases de données",
    "mysql",
    "mssql",
    "azure",
    "aws",
    "gcp",
    "data science",
    "marketing",
    
    
    # Compétences en Marketing liées à la Data
    "marketing"
    'analyse de marché',
    'segmentation de marché',
    'marketing automation',
    'analyse du comportement des utilisateurs',
    'analytique web',
    'optimisation des conversions',
    'a/b testing',
    'seo',
    'sea',
    'crm',
    'gestion de campagnes publicitaires',
    'analyse des performances marketing',
    'analyse de la concurrence',
    'intelligence concurrentielle',
    'analyse des tendances du marché',
    'stratégie de contenu',
    'analyse de sentiment',
    
    # Compétences en Développement
    'java',
    'javascript',
    'html',
    'css',
    'node.js',
    'react',
    'angular',
    'vue.js',
    'express.js',
    'django',
    'flask',
    'spring',
    'hibernate',
    'restful api',
    'graphql',
    
    # Compétences Talend
    'talend',
    'etl',
    'talend open studio',
    'talend data integration',
    'talend big data',
    'talend administration',
    
    # Compétences Snowflake
    'snowflake',
    'snowflake sql',
    'snowflake data warehousing',
    'snowflake administration',
    ]
    def extraire_competences(texte, competences):
        competences_trouvees = set()

        for competence in competences:
            # Utilisation d'une expression régulière pour rechercher la compétence dans le texte
            regex = re.compile(r'\b' + re.escape(competence) + r'\b', re.IGNORECASE)
            if regex.search(texte):
                competences_trouvees.add(competence)

        return competences_trouvees

# Appel de la fonction avec le texte et la liste des compétences
    competences_trouvees = extraire_competences(job_offer_text, competences)


# Affichage des compétences trouvées
    #print(competences_trouvees)
    
    def extraire_dates_et_calculer_experience(texte_cv):
        sections = re.split(r'd *i *p *l *[ôo] *m *e|f *o *r *m *a *t *i *o *n', texte_cv, flags=re.IGNORECASE)
        # Dictionnaire pour la conversion des noms de mois en numéros
        mois_en_numeros = {
            'janvier': 1, 'février': 2, 'mars': 3, 'avril': 4,
            'mai': 5, 'juin': 6, 'juillet': 7, 'août': 8,
            'septembre': 9, 'octobre': 10, 'novembre': 11, 'décembre': 12, 'Octobre':10
        }
        "[Dd](?:\'|e|epuis)? *"
        #print(re.findall(r'[Dd](?:epuis)? *(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4})', texte_cv))
        #print(re.findall(r'[Dd](?:\'|e)? *(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4}) +à +(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4})', texte_cv))
            # Modèle d'expression régulière pour rechercher les dates
        pattern_dates = re.compile(r'[Dd](?:\'|e)? *(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4}) +à +(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4})')
        pattern_date = re.compile(r'[Dd](?:epuis)? *(janvier|février|mars|avril|mai|juin|juillet|août|septembre|octobre|novembre|décembre) +(\d{4})')
        # Recherche des motifs dans le texte
        """matches = pattern_dates.findall(texte_cv)
        matches_app = pattern_date.findall(texte_cv)"""

        # Liste pour stocker les durées d'expérience en mois
        durees_experience = []
        for section in sections:
            #print(section.lower().replace(' ', ''))
            # Vérification si la section contient le mot-clé "expérience"
            
            if 'expérience'  in section.lower().replace(' ', ''):
                #print("-------------------condition ok-----------------------------------------")
                # Recherche des motifs dans la section
                #print(section)
                matches = pattern_dates.findall(section)
                matches_app = pattern_date.findall(section)
        # Parcours des correspondances
                for match in matches:
                    mois_debut, annee_debut, mois_fin, annee_fin = match

                            # Conversion du mois en numéro
                    mois_debut_numero = mois_en_numeros[mois_debut.lower()]
                    mois_fin_numero = mois_en_numeros[mois_fin.lower()]
                 #   print(mois_debut,mois_fin)
                  #  print(mois_debut_numero)
                   # print(mois_fin_numero)

                            # Ajout du jour par défaut pour éviter l'erreur de format
                    date_debut = datetime.strptime(f"1 {mois_debut_numero} {annee_debut}", "%d %m %Y")
                    date_fin = datetime.strptime(f"1 {mois_fin_numero} {annee_fin}", "%d %m %Y")
                  #  print(date_debut, date_fin)
                   # print((date_fin - date_debut).days / (30*12))
                            # Calcul de la durée en mois
                    durees_experience.append((date_fin - date_debut).days / (30*12))
                
                for match in matches_app:
                    mois_debut, annee_debut = match

                            # Conversion du mois en numéro
                    mois_debut_numero = mois_en_numeros[mois_debut.lower()]

                  #  print( mois_debut, mois_debut_numero)
                   # print((datetime.now() - datetime.strptime(f"1 {mois_debut_numero} {annee_debut}", "%d %m %Y")).days // (30*12))

                            # Ajout du jour par défaut pour éviter l'erreur de format
                    date_debut = datetime.strptime(f"1 {mois_debut_numero} {annee_debut}", "%d %m %Y")
                    durees_experience.append((datetime.now() - date_debut).days / (30*12))
        
                #print(durees_experience)
                # Calcul de la durée totale en mois
                duree_totale = sum(durees_experience)
                #print(duree_totale)
                # Conversion de la durée totale en années (arrondi à une décimale)
                duree_totale_en_annees = round(duree_totale, 1)

                return duree_totale_en_annees
    def extract_compétence(cv):
    # Traiter le CV et l'offre d'emploi avec spaCy
        doc_cv = nlp(cv)

    
        competences_cv = [ent.text.lower() for ent in doc_cv.ents if ent.label_ == 'compétence']
        return competences_cv
    
    def extract_name(text):
        doc = nlp(text)
        names = [ent.text for ent in doc.ents if ent.label_ == 'PER']
        if names:
            return names[0]
        else:
            return None
    """def extract_localisation(text):
        doc=nlp(text)
        locations = [ent.text for ent in doc.ents if ent.label_ == "LOC"]
        if locations:
            return locations[0]
        else:
            return None""" 
        # Use regular expression to find patterns that match names
        # name_pattern = r'\b[A-Z][a-z]+(?:\s[A-Z][a-z]+)*\b'
        # names = re.findall(name_pattern, text)
        # if names:
        #     return names[0]  # Assuming the first match is the name
        # else:
        #     return ""

    def extract_email(text):
        # Use regular expression to find email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,7}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            return emails[0]  # Assuming the first match is the email
        else:
            return ""

    def extract_phone(text):
        # Use regular expression to find phone numbers
        phone_pattern = r'\b(?:\+\d{1,2}\s?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        phones = re.findall(phone_pattern, text)
        if phones:
            return phones[0]  # Assuming the first match is the phone number
        else:
            return ""

    





    #cv_info = pd.DataFrame({"Email":[],"Score Correspandance %":[]})
    # Loop through CVs in the folder
    for cv_filename in os.listdir(cv_folder_path):
        if cv_filename.endswith('.pdf'):
            cv_path = os.path.join(cv_folder_path, cv_filename)
            cv_text = extract_text_from_pdf(cv_path)
        else:
            continue

        cv_text = preprocess_text(cv_text)
        #print(cv_text)
        comp = extraire_competences(cv_text, competences)
        exp = extraire_dates_et_calculer_experience(cv_text)
        name = extract_name(cv_text)
        email = extract_email(cv_text)
        phone = extract_phone(cv_text)
        #locations = extract_localisation(cv_text)
        # Initialize the TF-IDF vectorizer
        """vectorizer = TfidfVectorizer()

        # Transform the corpus into TF-IDF vectors
        tfidf_matrix = vectorizer.fit_transform([cv_text, job_offer_text])

        # Calculate the cosine similarity between the job offer and the CV
        cosine_similarity_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]"""
        vectorizer = TfidfVectorizer()
        tfidf_cvs = vectorizer.fit_transform(competences_trouvees)
        tfidf_job = vectorizer.transform(comp)
        cosine_similarities = cosine_similarity(tfidf_job, tfidf_cvs)
        model = SentenceTransformer('all-mpnet-base-v2')
        #print(list(comp))
        #print("['"+' '.join(list(comp))+"']")
        #print("['"+' '.join(list(competences_trouvees))+"']")
        cvs_embeddings = model.encode("[['"+' '.join(list(comp))+"']]", convert_to_tensor=True)
        #print(cvs_embeddings)
        cvs_embeddings = cvs_embeddings.reshape(1, -1)
        offre_embedding = model.encode("[['"+' '.join(list(competences_trouvees))+"']]", convert_to_tensor=True)
        #print(cvs_embeddings)
        offre_embedding = offre_embedding.reshape(1, -1)
        #print(offre_embedding)
        tensor1_normalized = F.normalize(cvs_embeddings, p=2, dim=1)
        tensor2_normalized = F.normalize(offre_embedding, p=2, dim=1)
        cosine_scores = F.cosine_similarity(tensor1_normalized, tensor2_normalized)
        cosine_scores = '{:.2f}'.format(cosine_scores.item())
        # Extract name, email, and phone from the CV
        # name = ""  # Implement logic to extract name
        # email = ""  # Implement logic to extract email
        # phone = ""  # Implement logic to extract phone

        # Create a dictionary with the extracted information
        cv_info = {
            "expérience" : exp,
            "comp": ' '.join(list(comp)),
            "email": email,
            #"phone": phone,
            "pourcentage_similarite": cosine_scores,
            #"location" : locations
        }

        

        cv_info_list.append(cv_info)
        #cv_info_list = cv_info.sort_values("Score Correspandance %", ascending=False)
    cv_info_list = sorted(cv_info_list, key=lambda x: x['pourcentage_similarite'], reverse=True)
    #print(cv_info_list)
    return print(cv_info_list)


traitement(r'C:\Users\guedrouz.ESH\Downloads\web_dev\dev_samrtCV\back\uploads\cv',r'C:\Users\guedrouz.ESH\Downloads\web_dev\dev_samrtCV\back\uploads\jobs\Data_Analyst.docx')

# 'cv_info_list' now contains the information for each CV in the folder, including similarity percentage with the job offer.
